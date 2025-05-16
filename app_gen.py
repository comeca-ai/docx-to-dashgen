import streamlit as ou até `10000`.
3.  **Simplifique a Amostra das Tabelas no st
from docx import Document
import pandas as pd
import plotly.express as px
import google.generativeai as gen Prompt:**
    ```python
    # Em analisar_documento_com_gemini
    sample_df = dfai
import json
import os
import traceback
import re 

# --- 1. Configuração da Ch.head(1).iloc[:, :min(3, len(df.columns))] # Apenas 1 linha,ave da API do Gemini ---
def get_gemini_api_key():
    try: 
        return st.secrets["GOOGLE_API_KEY"] # <<<<< CORRIGIDO: Apenas o nome da chave
    except (FileNotFoundError, KeyError): 
        api_key = os.environ.get("GOOGLE_API_KEY")
        return api_key if api_key else None
    except Exception as e: # Captura outros possíveis erros ao acessar secrets
        # st.error(f"Erro ao tentar acessar a chave da API do Gemini dos segredos: {e}") # Opcional: logar o erro
        # Tenta pegar da variável de ambiente como fallback se secrets falhar por outro motivo
        api_key = os.environ.get("GOOGLE_API_KEY")
        if api_key:
            # st.info("Chave da API do Gemini carregada da variável de ambiente após falha nos segredos.")
            return api_key
        # st.warning("Chave da API do Gemini não encontrada nos segredos nem nas variáveis de ambiente.") # Movido para a função que usa a chave
        return None

# --- 2. Fun 3 colunas de amostra
    colunas_para_mostrar_tipos = df.columns.tolist()[:min(3, len(df.columns))] # Apenas os 3 primeiros tipos de coluna
    ```
4.  **ções de Processamento do Documento e Interação com Gemini ---
def parse_value_for_numeric(val_str_in):
    if pd.isna(val_str_in) or str(val_str_in).strip() == '': return None
    text = str(val_str_in).strip()
    is_negative_paren = text.startswith('(') and text.endswith(')')
    if is_Teste com Conteúdo Mínimo:**
    *   Comente temporariamente a inclusão das tabelas no prompt (`tabelas_prompt_str = ""`). A API funciona só com o texto?
    *   Comente tempornegative_paren: text = text[1:-1]
    text_num_part = re.sub(r'[R$\s%€¥£]', '', text) 
    if ',' in text_num_part and '.'ariamente a inclusão do texto no prompt (`prompt_text = "Analise as tabelas abaixo."`). A API in text_num_part:
        if text_num_part.rfind('.') < text_num_part.r funciona só com as tabelas?
    Isso ajudará a isolar se o problema é o volume do texto,find(','): text_num_part = text_num_part.replace('.', '') 
        text_num_part = text_num_part.replace(',', '.') 
    elif ',' in text_num_part: text_num_part = text_num_part.replace(',', '.')
    match = re.search(r"([-+]?\d*\.?\d+|\d+)", text_num_part)
    if match: o volume/formato das tabelas, ou uma combinação.

**Versão Final Corrigida (Foco na Cons
        try: num = float(match.group(1)); return -num if is_negative_paren else num
        except ValueError: return None
    return None

def extrair_conteudo_docx(uploaded_istência do `session_state` e no `st.rerun()`):**

```python
import streamlit as st
from docx import Document
import pandas as pd
import plotly.express as px
import google.generativeai as genai
file):
    try:
        document = Document(uploaded_file)
        textos = [p.text for p in document.paragraphs if p.text.strip()]
        tabelas_data = [] 
        for i, table_obj in enumerate(document.tables):
            data_rows, keys, nome_tabela = [], None,import json
import os
import traceback
import re 

# --- 1. DEFINIÇÃO DE TODAS AS FUNÇÕES PRIMEIRO ---

def get_gemini_api_key():
    try: return st.secrets["GOOGLE_API_KEY f"Tabela_DOCX_{i+1}"
            try: 
                prev_el = table_obj._element.getprevious()
                if prev_el is not None and prev_el.tag.endswith('p'):
                    p_text = "".join(node.text for node in prev_el.xpath('.//w:t')).strip()
                    if p_text and len(p_text) < 8"]
    except (FileNotFoundError, KeyError): 
        api_key = os.environ.get("GOOGLE_API_KEY")
        return api_key if api_key else None

def parse_value_for_numeric(val_str_in):
    if pd.isna(val_str_in) or str(val_str_in).strip() == '': return None
    text = str(val_str_in).strip()
    is_negative_paren =0: nome_tabela = p_text.replace(":", "").strip()[:70] 
            except Exception: pass
            
            if len(table_obj.rows) > 0:
                header_cells = [cell.text.strip().replace("\n", " ") for cell in table_obj.rows[0].cells]
                keys = [key if key else f"Col{c_idx+1}" for c_idx, key in enumerate(header_cells)]
                for r_idx, row in enumerate( text.startswith('(') and text.endswith(')')
    if is_negative_paren: text = text[1:-1]
    text_num_part = re.sub(r'[R$\s%€¥£]', '', text) 
    if ',' in text_num_part and '.' in text_num_part:
        if text_num_part.rfind('.') < text_num_part.rfind(','): text_num_part = text_num_part.replace('.', '') 
        text_num_part = text_num_part.replace(',', '.') 
    elif ',' in text_num_part: text_num_part = text_num_part.replace(',',table_obj.rows):
                    if r_idx == 0: continue 
                    cells = [c.text.strip() for c in row.cells]
                    if keys: 
                        row_dict = {}
                        for k_idx, key_name in enumerate(keys):
                            row_dict[key_name] = cells[k_idx] if k_idx < len(cells) else None
                        data_rows.append(row_dict)
            if data_rows:
                try:
                    df = pd.DataFrame(data_rows)
                    for col in df.columns:
                        original_series = df[col].copy()
                        num_series = original_series '.')
    match = re.search(r"([-+]?\d*\.?\d+|\d+)", text_num_part)
    if match:
        try: num = float(match.group(1)); return -num if is_negative_paren else num
        except ValueError: return None
    return None

def extrair_conteudo_docx(uploaded_file):
    try:
        document = Document(uploaded_file)
        textos = [p.text for p in document.paragraphs if p.text.strip()]
        tabelas_data = [] 
        for i, table_obj in enumerate(document.tables):
            data_rows, keys, nome_tabela = [], None, f"Tabela_DOCX_{i+1}"
            try:.astype(str).apply(parse_value_for_numeric)
                        if num_series.notna().sum() / max(1, len(num_series)) >= 0.3: 
                            df[col] = pd.to_numeric(num_series, errors='coerce')
                            continue 
                        else: df[col] = original_series 
                        try:
                            temp_str_col = df[col].astype(str)
                            dt_series = pd.to_datetime(temp_str_col, errors='coerce', dayfirst=True) 
                            if dt_series.notna().sum() / max(1, len(dt_series)) >= 0.5: 
                                df[col] = dt_series
                            else: df[col] = original_series.astype(str).fillna('')
                        except Exception: df[col] = original_series 
                prev_el = table_obj._element.getprevious()
                if prev_el is not None and prev_el.tag.endswith('p'):
                    p_text = "".join(node.text for node in prev_el.xpath('.//w:t')).strip()
                    if p_text and len(p_text) < 80: nome_tabela = p_text.replace(":", "").strip()[:70] 
            except Exception: pass
            
            if len(table_obj.rows) > 0:
                header_cells = [cell.text.strip().replace("\n", " ") for cell in table_obj.rows[0].cells]
                keys = [key if key else f"Col{c_idx+1}" for c_idx, key in enumerate(header_cells)]
                for r_idx, row in enumerate(table_obj.rows):
                    if r_idx == 0: continue 
                    cells = [c.text.strip() for c in.astype(str).fillna('')
                    for col in df.columns: 
                        if df[col].dtype == 'object': df[col] = df[col].astype(str).fillna('')
                    tabelas_data.append({"id": f"doc_tabela_{i+1}", "nome": nome_tabela, "dataframe": df})
                except Exception as e_df_proc:
                    st.warning(f"Não foi possível processar DataFrame para tabela '{nome_tabela}': {e_df_proc}")
        return "\n\n".join(textos), tabelas_data
    except Exception as e_doc_read: 
        st.error(f"Erro crítico ao ler DOCX: {e_doc_read}")
        return "", []

def analisar_documento_com_gemini(texto_doc, tabelas_info_list):
    api_key = get_gemini_api_key()
    if not api_key: st.warning("Chave API Gemini não configurada."); return []
    try:
        genai.configure(api_key=api_key)
        safety_settings = [{"category": c,"threshold": "BLOCK_NONE"} for c in ["HARM_CATEGORY_HARASSMENT","HARM_ row.cells]
                    if keys: 
                        row_dict = {}
                        for k_idx, key_name in enumerate(keys):
                            row_dict[key_name] = cells[k_idx] if k_idx < len(cells) else None
                        data_rows.append(row_dict)
            if data_rows:
                try:
                    df = pd.DataFrame(data_rows)
                    for col in df.columns:
                        original_series = df[col].copy()
                        num_series = original_series.astype(str).apply(parse_value_for_numeric)
                        if num_series.notna().sum() / max(1, len(num_series)) >= 0.3: 
                            df[col] = pd.to_numeric(num_series, errors='coerce')
                            continue 
                        else: df[col] = original_series 
                        try:
                            temp_str_col = df[col].astype(str)
                            dt_series = pd.to_datetime(temp_str_col, errors='coerce', dayfirst=True) 
                            if dt_series.notna().sum() / max(1, len(dt_series)) >= 0.5: 
                                df[col] = dt_series
                            else: df[col] = original_series.astype(str).fillna('')
                        except Exception: df[col] = original_series.astype(str).CATEGORY_HATE_SPEECH","HARM_CATEGORY_SEXUALLY_EXPLICIT","HARM_CATEGORY_DANGEROUS_CONTENT"]]
        model = genai.GenerativeModel(model_name="gemini-1.5-flash-latest", safety_settings=safety_settings)
        tabelas_prompt_str = ""
        for t_info in tabelas_info_list:
            df, nome_t, id_t = t_info["dataframe"], t_info["nome"], t_info["id"]
            sample_df = df.head(2).iloc[:, :min(3, len(df.columns))] # Amostra ainda menor
            md_table = ""; 
            try: md_table = sample_df.to_markdown(index=False)
            except: md_table = sample_df.to_string(index=False) 
            colunas_para_mostrar_tipos = df.columns.tolist()[:min(5, len(df.columns))] # Menos colunas nos tipos
            col_types_list = [f"'{col_name_prompt}' (tipo: {str(df[col_name_prompt].dtype)})" for col_name_prompt in colunas_para_mostrar_tipos]
            col_types_str = ", ".join(col_types_list)
fillna('')
                    for col in df.columns: 
                        if df[col].dtype == 'object': df[col] = df[col].astype(str).fillna('')
                    tabelas_data.append({"id": f"doc_tabela_{i+1}", "nome": nome_tabela, "dataframe": df})
                except Exception as e_df_proc:
                    st.warning(f"Não foi possível processar DataFrame para tabela '{nome_tabela}': {e_df_proc}")
        return "\n\n".join(textos), tabelas_data
    except Exception as e_doc_read: 
        st.error(f"Erro crítico ao ler DOCX: {e_doc_read}")
        return "", []

def analisar_documento_com_gemini(texto_doc, tabelas_info_list):
    api_key = get_gemini_api_key()
    if not api_key: 
        st.warning("Chave API Gemini não configurada. Sugestões da IA desabilitadas.")
        return [] 
    try:
        genai.configure(api_key=api_key)
        safety_settings = [{"category": c,"threshold": "BLOCK_NONE"} for c in ["HARM_CATEGORY_HARASSMENT","HARM_CATEGORY_HATE_SPEECH","HARM_CATEGORY_SEXUALLY_EXPLICIT","HARM_CATEGORY_DANGEROUS_CONTENT"]]
        model = genai.GenerativeModel(model_name="gemini-1.5-flash-latest", safety_            tabelas_prompt_str += f"\n--- Tabela '{nome_t}' (ID: {id_t}) ---\nColunas e tipos (amostra): {col_types_str}\nAmostra de dados:\n{md_table}\n"
        text_limit = 20000 # Reduzido drasticamente para teste de erro 400
        prompt_text = texto_doc[:text_limit] + ("\n[TEXTO TRUNCADO...]" if len(texto_doc) > text_limit else "")
        prompt = f"""
        Você é um assistente de análise de dados. Analise o texto e as tabelas.
        [TEXTO]{prompt_text}[FIM TEXTO]
        [TABELAS]{tabelas_prompt_str}[FIM TABELAS]

        Gere lista JSON de sugestões. Objeto DEVE ter: "id", "titulo", "tipo_sugerido" ("kpi", "tabela_dados", "lista_swot", "grafico_barras", "grafico_pizza", "grafico_linha", "grafico_dispersao", "grafico_barras_agrupadas"), "fonte_id" (ID tabela ou "texto_desc_fonte"), "parametros" (objeto), "justificativa".
        Para "parametros":
        - "kpi": {{"valor": "ValorKPI", "delta": "Mudança", "descricao": "Contexto"}}
        - "tabela_dados": TABELA: {{"id_tabela_original": "ID_Tabela"}}. TEXTO: {{"dados": [{{"Col1": "V1"}}, ...], "colunas_titulo": ["Tsettings=safety_settings)
        tabelas_prompt_str = ""
        for t_info in tabelas_info_list:
            df, nome_t, id_t = t_info["dataframe"], t_info["nome"], t_info["id"]
            sample_df = df.head(1).iloc[:, :min(3, len(df.columns))] # Amostra BEM pequena
            md_table = ""
            try: md_table = sample_df.to_markdown(index=False)
            except: md_table = sample_df.to_string(index=False) 
            colunas_para_mostrar_tipos = df.columns.tolist()[:min(4, len(df.columns))] # Menos colunas
            col_types_list = [f"'{col_name_prompt}' (tipo: {str(df[col_name_prompt].dtype)})" for col_name_prompt in colunas_para_mostrar_tipos]
            col_types_str = ", ".join(col_types_list)
            tabelas_prompt_str += f"\n--- Tabela '{nome_t}' (ID: {id_t}) ---\nColunas e tipos (amostra): {col_types_str}\nAmostra dados:\n{md_table}\n"
        text_limit = 20000 # Limite de texto reduzido para teste de erro 400
        prompt_text = texto_doc[:text_limit] + ("\n[TEXTO TRUNCADO...]" if len(texto_doc) > text_limit else "")
        prompt = f"""
        Você é um assistente de análise de dados. Analise o texto e asCol1"]}}
        - "lista_swot": {{"forcas": ["F1"], "fraquezas": ["Fr1"], "oportunidades": ["Op1"], "ameacas": ["Am1"]}}
        - Gráficos TABELA ("barras", "linha", "dispersao"): {{"eixo_x": "NOME_COL_X", "eixo_y": "NOME_COL_Y"}} (Y numérico).
        - Gráficos PIZZA TABELA: {{"categorias": "NOME_COL_CAT", "valores": "NOME_COL_VAL_NUM"}} (Valores numéricos).
        - Gráficos DADOS TEXTO: {{"dados": [{{"EixoX": "A", "EixoY": 10.0}}, ...], "eixo_x": "EixoX", "eixo_y": "EixoY"}} (Valores NUMÉRICOS).
        - "grafico_barras_agrupadas": TABELA: {{"eixo_x": "COL_PRINC", "eixo_y": "COL_VAL", "cor_agrupamento": "COL_SUB"}}. DADOS TEXTO: {{"dados": [{{"CatP": "A", "SubC": "X", "Val": 10.0}}, ...], "eixo_x": "CatP", "eixo_y": "Val", "cor_agrupamento": "SubC"}}.
        CR tabelas.
        [TEXTO]{prompt_text}[FIM TEXTO]
        [TABELAS]{tabelas_prompt_str}[FIM TABELAS]
        Gere lista JSON de sugestões. Objeto: "id", "titulo", "tipo_sugerido" ("kpi", "tabela_dados", "lista_swot", "grafico_barras", "grafico_pizza", "grafico_linha", "grafico_dispersao", "grafico_barras_agrupadas"), "fonte_id", "parametros", "justificativa".
        "parametros":
        - "kpi": {{"valor": "ValorKPI", "delta": "Mudança", "descricao": "Contexto"}}
        - "tabela_dados": TABELA: {{"id_tabela_original": "ID_Tabela"}}. TEXTO: {{"dados": [], "colunas_titulo": []}}
        - "lista_swot": {{"forcas": [], "fraquezas": [], "oportunidades": [], "ameacas": []}}
        - Gráficos TABELA: {{"eixo_x": "COL_X", "eixo_y": "COL_Y"}} (Y numérico). Pizza: {{"categorias": "COL_CAT", "valores": "COL_VAL_NUM"}}.
        - Gráficos DADOS TEXTO: {{"dados": [], "eixo_x": "NomeX", "eixo_y": "NomeY"}} (Valores numéricos).
        - Barras Agrupadas TABÍTICO: Use NOMES EXATOS de colunas. Se valor não numérico, NÃO sugira gráfico numérico para ele A MENOS que extraia valor numérico (ex: '70%'->70.0). Se extrair, coloque em "dados" com VALORES NUMÉRICOS. SWOTs: individuais por player de tabela comparativa.
        Retorne APENAS a lista JSON.
        """
        with st.spinner("🤖 Gemini analisando..."):
            response = model.generate_content(prompt)
        cleaned_text = response.text.strip().lstrip("```json").rstrip("```").strip()
        sugestoes = json.loads(cleaned_text)
        if isinstance(sugestoes, list) and all(isinstance(item, dict) for item in sugestoes): st.success(f"{len(sugestoes)} sugestões!"); return sugestoes
        st.error("Resposta Gemini não é lista JSON."); return []
    except json.JSONDecodeError as e: st.error(f"Erro JSON Gemini: {e}"); st.code(response.text if 'response' in locals() else "N/A", language="ELA: {{"eixo_x": "COL_X", "eixo_y": "COL_Y_NUM", "cor_agrupamento": "COL_SUB_CAT"}}.
        CRÍTICO: Use NOMES EXATOS de colunas. Se valor não numérico, EXTRAIA NÚMERO para "dados" ou NÃO SUGIRA gráfico numérico. Cobertura Geográfica: "tabela_dados" com "dados". SWOT comparativo: "lista_swot" INDIVIDUAL por player. APENAS lista JSON.
        """
        with st.spinner("🤖 Gemini analisando..."):
            # st.text_area("Debug Prompt:", prompt, height=100) 
            response = model.generate_content(prompt)
        cleaned_text = response.text.strip().lstrip("```json").rstrip("```").strip()
        # st.text_area("Debug Resposta Gemini:", cleaned_text, height=100)
        sugestoes = json.loads(cleaned_text)
        if isinstance(sugestoes, list): st.success(f"{len(text"); return []
    except Exception as e: st.error(f"Erro API Gemini: {e}"); st.text(traceback.format_exc()); return []

def render_kpis(kpi_sugestoes):
    if kpi_sugestoes:
        num_kpis = len(kpi_sugestoes); kpi_cols = st.columns(min(num_kpis, 4)) 
        for i, kpi_sug in enumerate(kpi_sugestoes):
            with kpi_cols[i % min(num_kpis, 4)]:
                params=kpi_sug.get("parametros",{}); delta_val=str(params.get("delta",""))
                st.metric(label=kpi_sug.get("titulo","KPI"),value=str(params.get("valor","N/A")),delta=delta_val if delta_val else None,help=params.get("descricao"))
        st.divider()

def render_swot_card(sugestoes)} sugestões!"); return sugestoes
        st.error("Resposta Gemini não é lista JSON."); return []
    except json.JSONDecodeError as e: st.error(f"Erro JSON Gemini: {e}"); st.code(response.text if 'response' in locals() else "N/A", language="text"); return []
    except Exception as e: st.error(f"Erro API Gemini: {e}"); return []

# --- Funções de Renderização ---
def render_kpis(kpi_sugestoes):
    if kpi_sugestoes:
        kpi_cols=st.columns(min(len(kpi_sugestoes),4));i=0
        for kpi in kpi_sugestoes:
            with kpi_cols[i%min(len(kpi_sugestoes),4)]:
                p=kpi.get("parametros",{titulo_completo_swot, swot_data):
    st.subheader(f"{titulo_completo_swot}") 
    col1, col2 = st.columns(2)
    swot_map = {"forcas": ("Forças 💪", col1), "fraquezas": ("Fraquezas 📉", col1), 
                "oportunidades": ("Oportunidades 🚀", col2), "ameacas": ("Ameaças ⚠️", col2)}
    for key_swot_category, (header_swot_render, col_target_swot_render) in swot_map.items():
        with col_target_swot_render:
            st.markdown(f"##### {header_swot_render}")
            points_swot_render = swot_data.get(key_swot_category, ["N/A (info. não fornecida)"])
            if not points_swot_render or not isinstance(points_swot_render, list) or not all(isinstance(p_swot, str) for p_swot in points_swot_render): 
                });d=str(p.get("delta",""))
                st.metric(kpi.get("titulo","KPI"),str(p.get("valor","N/A")),d if d else None,help=p.get("descricao"));i+=1
        st.divider()

def render_swot_card(titulo, swot_data):
    st.subheader(titulo);c1,c2=st.columns(2)
    smap={"forcas":("Forças 💪",c1),"fraquezas":("Fraquezas 📉",c1),"oportunidades":("Oportunidades 🚀",c2),"ameacas":("Ameaças ⚠️",c2)}
    for k,(h,ct) in smap.items():
        with ct:st.markdown(f"##### {h}");points_swot_render = ["N/A (formato de dados incorreto)"]
            if not points_swot_render: points_swot_render = ["N/A"] 
            for item_swot_render in points_swot_render: 
                st.markdown(f"<div style='margin-bottom: 5px;'>- {item_swot_render}</div>", unsafe_allow_html=True)
    st.markdown("<hr style='margin-top: 10px; margin-bottom: 20px;'>", unsafe_allow_html=True)

def render_plotly_chart(item_config, df_plot_input):
    if df_plot_input is None: st.warning(f"Dados não para '{item_config.get('titulo','S/T')}'."); return False
    df_plot = df_plot_input.copy()
    tipo,titulo,params = item_config.get("tipo_sugerido"),item_config.get("titulo"),item_config.ps=swot_data.get(k,["N/A"])
        if not ps or not isinstance(ps,list) or not all(isinstance(p_s,str) for p_s in ps):ps=["N/A (formato erro)"]
        if not ps:ps=["N/A"];[st.markdown(f"- {p_s}") for p_s in ps] # Removido div e key
    st.markdown("<hr style='margin-top:10px;margin-bottom:20px;'>",unsafe_allow_html=True)

def render_plotly_chart(item, df_in):
    if df_in is None:st.warning(f"Dadosget("parametros",{})
    x,y,cat,val,cor_grp = params.get("eixo_x"),params.get("eixo_y"),params.get("categorias"),params.get("valores"),params.get("cor_agrupamento")
    fig,fn,p_args = None,None,{}
    if tipo in ["grafico_barras","grafico_barras_agrupadas"] and x and y: fn,p_args=px.bar,{"x":x,"y":y}; 
        if tipo=="grafico_barras_agrupadas" and cor_grp:p_args["color"],p_args["barmode"]=cor_grp,"group"
    elif tipo=="grafico_linha" and x and y: fn,p_args=px.line,{"x":x,"y":y,"markers":True}
    elif tipo=="grafico_dispersao" and x and y: fn,p_args= não para '{item.get('titulo','S/T')}'.");return False
    df=df_in.copy();tipo,tit,prms=item.get("tipo_sugerido"),item.get("titulo"),item.get("parametros",{})
    x,y,cat,val,c_grp=prms.get("eixo_x"),prms.get("eixo_y"),prms.get("categorias"),prms.get("valores"),prms.get("cor_agrupamento")
    fn,p_args=None,{}
    if tipo in ["grafico_barras","grafico_barras_agrupadas"] and x and y:fn,p_args=px.bar,{"x":xpx.scatter,{"x":x,"y":y}
    elif tipo=="grafico_pizza" and cat and val: fn,p_args=px.pie,{"names":cat,"values":val}
    if fn:
        req_cols=[c for c in p_args.values() if isinstance(c,str)]
        if not all(c in df_plot.columns for c in req_cols): st.warning(f"Colunas {req_cols} não em '{titulo}'. Em DF: {df_plot.columns.tolist()}"); return False
        try:
            df_cln=df_plot.copy(); y_ax,val_ax=p_args.get("y"),p_args.get("values")
            if y_,"y":y};
                                                                      if tipo=="grafico_barras_agrupadas" and c_grp:p_args["color"],p_args["barmode"]=c_grp,"group"
    elif tipo=="grafico_linha" and x and y:fn,p_args=px.line,{"x":x,"y":y,"markers":True}
    elif tipo=="grafico_dispersao" and x and y:fn,p_args=px.scatter,{"x":x,"y":y}
    elif tipo=="grafico_pizza" and cat and val:fn,p_args=px.pie,{"names":cat,"values":val}
    if fn:ax and y_ax in df_cln.columns:df_cln[y_ax]=pd.to_numeric(df_cln[y_ax],errors='coerce')
            if val_ax and val_ax in df_cln.columns:df_cln[val_ax]=pd.to_numeric(df_cln[val_ax],errors='coerce')
            cols_chk_na=[v for v in p_args.values() if isinstance(v,str) and v in df_cln.columns]
            df_cln.dropna(subset=cols_chk_na,inplace=True)
            if not df_cln.empty: fig=fn(df_cln,title=titulo,**p_args);st.plotly_chart(fig,use_container_width=True);return
        req=[c for c in p_args.values() if isinstance(c,str)]
        if not all(c in df.columns for c in req):st.warning(f"Colunas {req} não em '{tit}'. Em DF:{df.columns.tolist()}");return False
        try:
            df_c=df.copy();y_a,v_a=p_args.get("y"),p_args.get("values")
            if y_a and y_a in df_c.columns:df_c[y_a]=pd.to_numeric(df_c[y_a],errors='coerce')
            if v_a and v_a in df_c.columns:df_c[v_a]=pd.to_numeric(df_c[v_a],errors='coerce') True
            else: st.warning(f"Dados insuficientes para '{titulo}' após limpar NaNs de {cols_chk_na}.")
        except Exception as e: st.warning(f"Erro Plotly '{titulo}': {e}. Dtypes: {df_plot.dtypes.to_dict() if df_plot is not None else 'DF None'}")
    elif tipo in ["grafico_barras","grafico_barras_agrupadas","grafico_linha","grafico_dispersao","grafico_pizza","grafico_radar"]:
        st.warning(
            df_c.dropna(subset=[c for c in p_args.values() if isinstance(c,str) and c in df_c.columns],inplace=True)
            if not df_c.empty:st.plotly_chart(fn(df_c,title=tit,**p_args),use_container_width=True);return True
            else:st.warning(f"Dados insuficientes para '{tit}' após NaNs.")
        except Exception as e:st.warning(f"Erro Plotly '{tit}':{e}. Dtypes:{df.dtypes.to_dict() if dff"Params incompletos para '{titulo}' ({tipo}).")
    return False

# --- 3. Interface Streamlit Principal ---
st.set_page_config(layout="wide", page_title="Gemini DOCX Insights GEN")
for k, dv in [("sugestoes_gemini",[]),("config_sugestoes",{}),("conteudo_docx",{"texto":"","tabelas":[]}),
              ("nome_arquivo_atual",None),("debug_checkbox_key",False),("pagina_selecionada","Dashboard Principal")]:
    st.session_state.setdefault is not None else 'DF None'}")
    elif tipo in ["grafico_barras","grafico_barras_agrupadas","grafico_linha","grafico_dispersao","grafico_pizza"]:st.warning(f"Params incompletos: '{tit}' ({tipo}).")
    return False

# --- 3. Interface Streamlit Principal ---
st.set_page_config(layout="wide",page_title="Gemini DOCX Insights GEN")
for k,dv in [(k, dv)

st.sidebar.title("✨ Navegação"); pg_opts_sb=["Dashboard Principal","Análise SWOT Detalhada"]
st.session_state.pagina_selecionada=st.sidebar.radio("Selecione:",pg_opts_sb,index=pg_opts_sb.index(st.session_state.pagina_selecionada),key="nav_radio_gen_v9")
st.sidebar.divider(); uploaded_file_sb=st.sidebar.file_uploader("Selecione DOCX",type="docx",key="uploader_sb_gen_v9")
st("sugs_gemini",[]),("cfg_sugs",{}),("doc_ctx",{"texto":"","tabelas":[]}),
              ("nome_arquivo_atual",None),("dbg_cb_key",False),("pg_sel","Dashboard Principal")]:st.session_state.setdefault(k,dv)

st.sidebar.title("✨ Navegação");pg_opts=st.session_state.pg_sel_.session_state.debug_checkbox_key=st.sidebar.checkbox("Mostrar Debug Info",value=st.session_state.debug_checkbox_key,key="debug_cb_gen_v9")

if uploaded_file_sb:
    if st.session_state.nome_arquivo_atual!=uploaded_file_sb.name: 
        with st.spinner("Processando novo documento..."):
            st.session_state.sugestoes_gemini,st.sessionopts if "pg_sel_opts" in st.session_state else ["Dashboard Principal","Análise SWOT Detalhada"]
st.session_state.pg_sel=st.sidebar.radio("Selecione:",_state.config_sugestoes=[],{}
            st.session_state.nome_arquivo_atual=uploaded_file_sb.name
            txt_main,tbls_main=extrair_conteudo_docx(uploaded_file_sb);st.session_state.conteudo_docx={"texto":txt_main,"tabelas":tblspg_opts,index=pg_opts.index(st.session_state.pg_sel),key="nav_radio_k")
st.session_state.pg_sel_opts=pg_opts # Guarda as opções para consistência
st.sidebar.divider();up_file=st.sidebar.file_uploader("Selecione_main}
            if txt_main or tbls_main:
                sugs_main=analisar_documento_com_gemini(txt_main,tbls_main);st.session_state.sugestoes_gemini=sugs_main
                st.session_state.config_sugestoes={s.get("id",f"s_{i}_{hash(s.get('titulo'))}"):{"aceito":True,"titulo_editado":s DOCX",type="docx",key="upl_k")
st.session_state.dbg_cb_key=st.sidebar.checkbox("Debug Info",value=st.session_state.dbg_cb_key,key="dbg_k")

if up.get("titulo","S/T"),"dados_originais":s} for i,s in enumerate(sugs_main)}
            else: st.sidebar.warning("Nenhum conteúdo extraído.")
    if st.session_state.debug_checkbox_key and (st.session_state.conteudo_docx["texto"] or st.session_state.conteudo_docx["tabelas_file:
    if st.session_state.nome_arquivo_atual!=up_file.name: 
        with st.spinner("Processando..."):
            st.session_state.sugs_gemini,st.session_state.cfg_sugs=[],{}
            st.session_state.nome_arquivo_atual=up_file.name
"]):
        with st.expander("Debug: Conteúdo DOCX (após extração e tipos)",expanded=False):
            st.text_area("Texto (amostra)",st.session_state.conteudo_docx["texto"][:1000],height=80)
            for t_dbg in st.session_state.conteudo_docx["tabelas"]:            txt,tbls=extrair_conteudo_docx(up_file);st.session_state.doc_ctx={"texto":txt,"tabelas":tbls}
            if txt or tbls:
                sugs=analisar_documento_com_gemini(txt,tbls);st.session_state.sugs_gemini=sugs

                st.write(f"ID: {t_dbg['id']}, Nome: {t_dbg['nome']}")
                try:st.dataframe(t_dbg['dataframe'].head().astype(str).fillna("-"))
                except:st.text(f"Head:\n{t_dbg['dataframe'].head().to_string(na_rep='-')}")
                st.write("Tipos:",t_dbg['dataframe'].dtypes.to_dict())
                st.session_state.cfg_sugs={s.get("id",f"s_{i}_{hash(s.get('titulo'))}"):{"aceito":True,"tit_edit":s.get("titulo","S/T"),"orig_data":s} for i,s in enumerate(sugs)}
            else:st.sidebar.warning("Nenhum conteúdo extraído.")
    if st.session_state.dbg_cb_key and (st.session_state.doc_ctx["texto"] or st.session_state.doc_ctx["tabelas"]):
        with st.exp    if st.session_state.sugestoes_gemini:
        st.sidebar.divider();st.sidebar.header("⚙️ Configurar Sugestões")
        for sug_cfg in st.session_state.sugestoes_gemini:
            s_id,cfg=sug_cfg.get('id'),st.session_state.config_sugestoes.get(sug_cfg.get('id'))
            if not s_id or not cfg:continue
            with st.sidebar.expander(f"{cfg['titulo_editado']}",expanded=False):
                st.caption(f"Tipo:{sug_cfg.get('tipoander("Debug: Conteúdo DOCX",expanded=False):
            st.text_area("Texto (amostra)",st.session_state.doc_ctx["texto"][:500],height=60)
            for t_dbg in st.session_state.doc_ctx["tabelas"]:
                st.write(f"ID:{t_dbg['id']},Nome:{t_dbg['nome']}");
                try:st.dataframe(t_dbg['dataframe'].head_sugerido')}|Fonte:{sug_cfg.get('fonte_id')}")
                cfg["aceito"]=st.checkbox("Incluir?",value=cfg["aceito"],key=f"acc_gen_{s_id}")
                cfg["titulo_editado"]=st.text_input("Título",value=cfg["titulo_editado"],key=f"tit_gen_{s_id}")
else: 
    if st.session_state.pagina().astype(str).fillna("-"))
                except:st.text(f"Head:\n{t_dbg['dataframe'].head().to_string(na_rep='-')}")
                st.write("Tipos:",t_dbg['dataframe'].dtypes.to_dict())
    if st.session_state.sugs_gemini:
        st.sidebar.divider();st.sidebar.header("⚙️ Configurar Sugestões")
        for sug_sb_selecionada=="Dashboard Principal":st.info("Upload DOCX na barra lateral.")

if st.session_state.pagina_selecionada=="Dashboard Principal":
    st.title("📊 Dashboard de Insights Genérico")
    if uploaded_file_sb and st.session_state.sugestoes_gemini:
        kpis,outros=[],[]
        for s_id,s_cfg in st.session_state.config_sugestoes.items():
            if s_cfg["aceito"]:item={"titulo":s_cfg["titulo_editado"], in st.session_state.sugs_gemini:
            s_id,cfg=sug_sb['id'],st.session_state.cfg_sugs.get(sug_sb['id'])
            if not cfg:cfg=st.session_state.cfg_sugs[s_id]={"aceito":True,"tit_edit":sug_sb.get("titulo","S/T"),"orig_data":sug_sb}
            with st.sidebar.expander(f"{cfg['tit_edit']}",expanded=False):
                st.caption(f"Tipo:{sug_**s_cfg["dados_originais"]};(kpis if item.get("tipo_sugerido")=="kpi" else outros).append(item)
        render_kpis(kpis)
        if st.session_state.debug_checkbox_key:
             with st.expander("Debug: Elementos para Dashboard (Não-KPI)",expanded=True):st.json({"Outros":outros},expanded=False)
        
        elements_renderedsb.get('tipo_sugerido')}|Fonte:{sug_sb.get('fonte_id')}")
                cfg["aceito"]=st.checkbox("Incluir?",value=cfg["aceito"],key=f"ac_{s_id}")
                cfg["tit_edit"]=st.text_input("Título",value=cfg["tit_edit"],key=f"ti_{s_id}")
else: 
    if st.session_state.pg_sel=="Dashboard Principal":st.info("Upload DOCX na barra lateral.")

if st.session_state.pg__count = 0 # Inicialização CORRETA
        col_render_idx = 0 # Inicialização CORRETA

        if outros:
            item_cols_main_render=st.columns(2)
            for item_loop in outros:
                if item_loop.get("tipo_sugerido")=="lista_swot":sel=="Dashboard Principal":
    st.title("📊 Dashboard de Insights");
    if up_file and st.session_state.sugs_gemini:
        kpis,outros=[],[];[(kpis if s_c["origcontinue
                
                rendered_this_item = False # Flag para esta iteração
                with item_cols_main_render[col_render_idx % 2]:
                    st.subheader(item_loop["titulo"]);df_plot__data"].get("tipo_sugerido")=="kpi" else outros).append({"titulo":s_c["tit_edit"],**s_c["orig_data"]}) for s,s_c in st.session_state.item=None
                    params_item,tipo_item,fonte_item=item_loop.get("parametcfg_sugs.items() if s_c["aceito"]]
        render_kpis(kpisros",{}),item_loop.get("tipo_sugerido"),item_loop.get("fonte_id")
                    try:
)
        if st.session_state.dbg_cb_key:
             with st.expander("Debug                        if params_item.get("dados"):df_plot_item=pd.DataFrame(params_item[": Elementos para Dashboard (Não-KPI)",expanded=False):st.json({"Outros":outros})dados"])
                        elif str(fonte_item).startswith("doc_tabela_"):df_plot_item=next((t["dataframe"] for t in st.session_state.doc_ctx["tabelas"] if t["id"]==fonte
        if outros:
            cols_d,idx_d,cnt_d=st.columns(2),0,0_item),None)
                        
                        if tipo_item=="tabela_dados":
                            df_tbl
            for item in outros:
                if item.get("tipo_sugerido")=="lista_swot":_item=None
                            if str(fonte_item).startswith("texto_") and params_item.get("dados"):
                                df_tbl_item=pd.DataFrame(params_item.get("dados"));
                                if params_item.get("colunas_continue
                with cols_d[idx_d%2]:
                    st.subheader(item["titulo"]);titulo"):df_tbl_item.columns=params_item.get("colunas_titulo")
                            else:id_tbl_item=params_item.get("id_tabela_original",fonte_item);df_tbl_itemdf_p,rend=None,False
                    prms,tipo,fonte=item.get("parametros",{}),item.get("tipo_sugerido"),item.get("fonte_id")
                    try:
                        if prms=next((t["dataframe"] for t in st.session_state.doc_ctx["tabelas"] if t["id"]==id_tbl_item),None)
                            if df_tbl_item is not None:try:st.dataframe(df_tbl_item.astype(str).fillna("-"));rendered_this_item=True
                               .get("dados"):df_p=pd.DataFrame(prms["dados"])
                        elif str(fonte).startswith("doc_tabela_"):df_p=next((t["dataframe"] for t in st.session_state.doc_ctx["tabelas"] if t["id"]==fonte),None)
                        if                 except:st.text(df_tbl_item.to_string(na_rep='-'));rendered_this_item=True # Considera renderizado
                            else:st.warning(f"Tabela '{item_loop['titulo']}' ( tipo=="tabela_dados":
                            df_t=None
                            if str(fonte).startswith("texto_") and prms.get("dados"):
                                df_t=pd.DataFrame(prms.get("dados"));
                                if prms.get("colunas_titulo"):df_t.columns=prms.get("colunas_titulo")
                            elseFonte:{fonte_item}) não encontrada.")
                        elif tipo_item in ["grafico_barras","grafico_linha","grafico_dispersao","grafico_pizza","grafico_barras_agrupadas"]:
                            if render:id_t=prms.get("id_tabela_original",fonte);df_t=next((t["dataframe"] for t in st.session_state.doc_ctx["tabelas"] if t["id"]==id_t),None)
                            if df_t is not None:try:st.dataframe(_plotly_chart(item_loop,df_plot_item):rendered_this_item=True
                        elif tipo_item=='mapa':st.info(f"Mapa '{item_loop['titulo']}' não implementado.");rendered_this_item=True
                        if not rendered_this_item and tipo_item not in ["kdf_t.astype(str).fillna("-"))
                                                except:st.text(df_t.to_string(na_rep='-'));rend=True
                            else:st.warning(f"Tabela '{item['titulo']}' (Fonte:{fonte}) não encontrada.")
                        elif df_p is not None and tipo in ["grafico_barras","grafpi","lista_swot","mapa"]:st.info(f"'{item_loop['titulo']}' ({tipo_item}) não gerado.")
                    except Exception as e:st.error(f"Erro render '{item_loop['titulo']}': {e}")
                if rendered_this_item:col_render_idx+=1;elements_rendered_count+=1 #ico_linha","grafico_dispersao","grafico_pizza","grafico_barras_agrupadas"]:
                            if render_plotly_chart(item,df_p):rend=True
                        elif tipo=='mapa':st.info(f"Mapa '{item['titulo']}' não implementado.");rend=True
 Usa a variável correta
            if elements_rendered_count==0 and any(c['aceito'] and c['dados_originais'].get('tipo_sugerido') not in ['kpi','lista_swot'] for c in st.session_state.cfg_sugs.values()):
                st.info("Nenhum                        if not rend and tipo not in["kpi","lista_swot","mapa"]:st.info(f"'{item['titulo']}' ({tipo}) não gerado.")
                    except Exception as e:st.error(f"Erro render '{item['titulo']}': {e}")
                if rend:idx_d+=1;cnt_d+=1
            if cnt_d==0 and any(c['aceito'] and c['orig_data'].get('tipo_sugerido') not in ['kpi','lista_swot'] for c in st.session_ gráfico/tabela (além de KPIs/SWOTs) pôde ser gerado.")
        elif not kpis and not uploaded_file_sb:pass
        elif not kpis and not outros and uploaded_file_sb and st.session_state.s_gemini:st.info("Nenhum elemento selecionado/gerado.")
elif st.session_state.pg_sel=="Análise SWOT Detalhada":
    st.title("🔬 Análise SWOT Detalhada")state.cfg_sugs.values()):
                st.info("Nenhum gráfico/tabela (além de KPIs/SWOTs) para Dashboard.")
        elif not kpis and not up_file:pass
        elif not kpis and not outros and up_file and st.session_state.sugs_gemini:st.info("Nenhum elemento selecionado/gerado.")
elif st.session_state.pg_sel=="Análise SWOT Detalhada":
    if not uploaded_file_sb:st.warning("Upload DOCX na barra lateral.")
    elif not st.session_state.s_gemini:st.info("Aguardando processamento/sugestões.")
    else:
        swot_sugs_pg=[s_cfg["dados_originais"] for s_id,s_cfg in st.session_state.cfg_sugs.items() if s_cfg["aceito"] and s_cfg["dados_originais"].
    st.title("🔬 Análise SWOT Detalhada")
    if not up_file:st.warning("Upload DOCX na barra lateral.")
    elif not st.session_state.sugs_gemini:st.info("Aguardando processamento/sugestões.")
    else:
        swot_s=[get("tipo_sugerido")=="lista_swot"]
        if not swot_sugs_pg:st.info("Nenhuma análise SWOT sugerida/selecionada.")
        else:
            if st.session_state.dbg_cb_key:
                with st.expander("Debug: Dados para Análise SWOT",expanded=False):st.json({"SWOTs":swot_sugs_pg})
            for sws_c["orig_data"] for s,s_c in st.session_state.cfg_sugs.items() if s_c["aceito"] and s_c["orig_data"].get("tipo_sugerido")=="lista_swot"]
        if not swot_s:st.info("Nenhuma análise SWOT selecionada.")
        else:
            if st.session_state.dbg_cb_key:
                with st.expander("Debug: Dados para SWOTot_item_pg in swot_sugs_pg:render_swot_card(swot_item_pg.get("titulo","SWOT"),swot_item_pg.get("parametros",{}),card_key_prefix=swot_item_pg.get("id","swot_pg"))
if uploaded_file_sb is None and st.session_state. (Página Dedicada)",expanded=False):st.json({"SWOTs":swot_s})
            for swot_i in swot_s:render_swot_card(swot_i.get("titulo","nome_arquivo_atual is not None:
    keys_to_preserve_cl=[k for k in st.session_state.keys() if k.startswith(("nav_radio_gen_v9","uploader_sb_gen_v9","SWOT"),swot_i.get("parametros",{}),swot_i.get("id","swot_pg"))
if up_file is None and st.session_state.nome_arquivo_atual is not None: # Corrigido paradebug_cb_gen_v9","acc_gen_","tit_gen_"))]
    curr_keys_cl= nome_arquivo_atual
    keys_to_preserve_on_clear = [k for k in st.session_state.keys() if k.endswith(("_key", "_widget_key")) or k.startswith("nav_radio_") or k.startswithlist(st.session_state.keys())
    for k_cl in curr_keys_cl:
        if k_cl not in keys_to_preserve_cl:
            if k_cl in st.session_state:del st.session("uploader_sidebar_") or k.startswith("debug_cb_")]
    current_keys_on_clear__state[k_cl]
    for k_rinit,dv_rinit in [("sugmain = list(st.session_state.keys())
    for key_cl_remove in current_keys_on_clear_mainestoes_gemini",[]),("config_sugestoes",{}),("conteudo_docx",{"texto":"",":
        if key_cl_remove not in keys_to_preserve_on_clear and not any(key_cl_removetabelas":[]}),
                             ("nome_arquivo_atual",None),("debug_checkbox_key",False),("pagina_selecionada","Dashboard Principal")]:
        st.session_state.setdefault(k_r.startswith(prefix) for prefix in ["acc_loop_gen_", "tit_loop_gen_"]):
            if key_cl_remove in st.session_state: del st.session_state[key_cl_removeinit,dv_rinit)
    st.rerun()
